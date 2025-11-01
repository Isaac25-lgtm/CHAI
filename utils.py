import os
import tempfile
import uuid
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

from config import Config
from logger import log_function_call, log_email_event
from validators import DataValidator

class ExcelGenerator:
    """Handle Excel file generation with proper error handling"""
    
    @staticmethod
    @log_function_call
    def create_participant_excel(participants_data: List[Dict], protect_sheet: bool = False) -> Tuple[str, str]:
        """Generate Excel file for participant registration
        
        Args:
            participants_data: List of participant dictionaries
            protect_sheet: If True, protect the sheet to prevent editing (for regular users)
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Participant Registration"
            
            headers = [
                "No.", "Participant's Name", "Cadre", "Duty Station (Facility)",
                "District", "Mobile Number Registered", 
                "Names Registered on Mobile Money (First & Last Names)"
            ]
            
            # Apply header styling
            ExcelGenerator._apply_header_styling(ws, headers)
            
            # Set column widths
            ExcelGenerator._set_column_widths(ws)
            
            # Write participant data
            ExcelGenerator._write_participant_data(ws, participants_data)
            
            # Generate filename and save
            filename = f'participant_registration_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            filepath = ExcelGenerator._save_workbook(wb, filename)
            
            return filepath, filename
            
        except Exception as e:
            raise Exception(f"Failed to create participant Excel file: {str(e)}")
    
    @staticmethod
    @log_function_call
    def create_assessment_excel(assessment_data: Dict, protect_sheet: bool = False) -> Tuple[str, str]:
        """Generate comprehensive Excel assessment report
        
        Args:
            assessment_data: Dictionary containing assessment information
            protect_sheet: If True, protect all sheets to prevent editing (for regular users)
        """
        try:
            wb = Workbook()
            
            # Detect if this is a section-specific download
            scores = assessment_data.get('scores', {})
            section_definitions = assessment_data.get('section_definitions', {})
            # Use ALL section keys from section_definitions instead of hardcoded list
            section_ids = list(section_definitions.keys())
            sections_with_data = [key for key in section_ids if key in scores or any(k.startswith(key + '_') for k in scores.keys())]
            is_section_download = len(sections_with_data) == 1
            
            # Create summary dashboard
            ExcelGenerator._create_summary_sheet(wb, assessment_data)
            
            # Create detailed scores sheet (filtered for section-specific downloads)
            ExcelGenerator._create_details_sheet(wb, assessment_data, is_section_download)
            
            # Create action plan sheet
            ExcelGenerator._create_action_plan_sheet(wb, assessment_data)
            
            # Generate filename and save
            facility_name = assessment_data.get('facilityName', 'Unknown').replace(' ', '_')
            filename = f'assessment_report_{facility_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            filepath = ExcelGenerator._save_workbook(wb, filename)
            
            return filepath, filename
            
        except Exception as e:
            raise Exception(f"Failed to create assessment Excel file: {str(e)}")
    
    @staticmethod
    def _apply_header_styling(ws, headers: List[str]):
        """Apply consistent header styling"""
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border
    
    @staticmethod
    def _set_column_widths(ws):
        """Set appropriate column widths"""
        column_widths = {
            'A': 6, 'B': 25, 'C': 20, 'D': 30, 
            'E': 20, 'F': 25, 'G': 35
        }
        
        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width
    
    @staticmethod
    def _write_participant_data(ws, participants_data: List[Dict]):
        """Write participant data to worksheet"""
        data_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        for idx, participant in enumerate(participants_data, 1):
            row_num = idx + 1
            
            # Sanitize and write data
            ws.cell(row=row_num, column=1).value = idx
            ws.cell(row=row_num, column=2).value = DataValidator.sanitize_input(participant.get('participantName', ''))
            ws.cell(row=row_num, column=3).value = DataValidator.sanitize_input(participant.get('cadre', ''))
            ws.cell(row=row_num, column=4).value = DataValidator.sanitize_input(participant.get('dutyStation', ''))
            ws.cell(row=row_num, column=5).value = DataValidator.sanitize_input(participant.get('district', ''))
            ws.cell(row=row_num, column=6).value = DataValidator.sanitize_input(participant.get('mobileNumber', ''))
            ws.cell(row=row_num, column=7).value = DataValidator.sanitize_input(participant.get('mobileMoneyName', ''))
            
            # Apply styling
            for col_num in range(1, 8):
                cell = ws.cell(row=row_num, column=col_num)
                cell.border = border
                cell.fill = data_fill
                cell.alignment = Alignment(vertical='center')
    
    @staticmethod
    def _create_summary_sheet(wb: Workbook, assessment_data: Dict):
        """Create comprehensive summary dashboard sheet"""
        ws = wb.active
        ws.title = "Assessment Summary"
        
        # Define borders and fills
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        info_fill = PatternFill(start_color="E8F4F8", end_color="E8F4F8", fill_type="solid")
        
        # Detect if this is a section-specific download
        scores = assessment_data.get('scores', {})
        section_definitions = assessment_data.get('section_definitions', {})
        # Use ALL section keys from section_definitions instead of hardcoded list
        section_ids = list(section_definitions.keys())
        
        sections_with_data = [key for key in section_ids if key in scores or any(k.startswith(key + '_') for k in scores.keys())]
        is_section_download = len(sections_with_data) == 1
        
        # Section name mapping
        section_name_map = {
            'triple_elimination_treatment': 'TRIPLE ELIMINATION (HIV, Syphilis and Hep B): Linkage to treatment',
            'art_pmtct': 'ART in PMTCT Facilities (VL uptake at ANC1 for ARTK)',
            'quality_pmtct': 'Quality of PMTCT services (Review 10 client\'s charts)',
            'patient_tracking': 'Patient Tracking HIV+ Pregnant Women',
            'adherence_support': 'Adherence Support',
            'facility_linkage': 'Facility Linkage to Community Care and Support Services for Adult PLHIV',
            'sti_screening': 'STI Screening and Management in HIV Clinics Serving General Population',
            'early_infant_diagnosis': 'Early Infant Diagnosis [HEI]',
            'ctx_hei': 'CTX for HIV-Exposed Infants [HEI]',
            'tracking_hei': 'Tracking HIV-Exposed Infants [HEI]',
            'enrolment_eid_art': 'Enrolment of HIV-Infected Infants (EID Services into ART Services)',
            'hei_eid_registers': 'HIV Exposed Infant/Early Infant Diagnosis Registers',
            'supply_chain_eid': 'Supply Chain Reliability (Early Infant Diagnosis) [HEI]',
            'supply_chain_pmtct': 'Supply Chain Reliability – HIV PMTCT (Mother)',
            'supply_chain_syphilis': 'Supply Chain Reliability – Syphilis PMTCT',
            'supply_chain_hepb': 'Supply Chain Reliability – Hepatitis B PMTCT (Maternal and Birth Dose)',
            'patient_records': 'Patient/Beneficiary Records'
        }
        
        # Calculate scores
        total_possible = 0
        total_scored = 0
        scored_indicators = 0
        section_name = None
        
        if is_section_download:
            section_key = sections_with_data[0]
            section_name = section_name_map.get(section_key, section_key.upper())
            total_possible = 4
            score = scores.get(section_key, 0)
            try:
                total_scored = float(score) if score not in ['N/A', '', None] else 0
            except (ValueError, TypeError):
                total_scored = 0
            scored_indicators = 1 if total_scored > 0 else 0
        else:
            for section_id in section_ids:
                total_possible += 4
                score = scores.get(section_id, 0)
                try:
                    score_val = float(score) if score not in ['N/A', '', None] else 0
                    if score_val > 0:
                        total_scored += score_val
                        scored_indicators += 1
                except (ValueError, TypeError):
                    pass
        
        percentage = round((total_scored / total_possible * 100), 1) if total_possible > 0 else 0
        
        # Determine performance level
        if percentage >= 75:
            performance = "EXCELLENT"
            perf_color = "006400"
        elif percentage >= 50:
            performance = "GOOD"
            perf_color = "90EE90"
        elif percentage >= 25:
            performance = "NEEDS IMPROVEMENT"
            perf_color = "FFC107"
        else:
            performance = "CRITICAL"
            perf_color = "DC3545"
        
        # Set column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 20
        
        # Add title
        row = 1
        ws.merge_cells(f'A{row}:C{row}')
        title_cell = ws[f'A{row}']
        if is_section_download:
            title_cell.value = f"ASSESSMENT SUMMARY\n{section_name}"
        else:
            title_cell.value = "PMTCT ASSESSMENT SUMMARY REPORT"
        title_cell.font = Font(size=16 if is_section_download else 18, bold=True, color="FFFFFF")
        title_cell.fill = header_fill
        title_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        title_cell.border = thin_border
        ws.row_dimensions[row].height = 40 if is_section_download else 30
        
        # Add facility information
        row += 2
        facility_info = [
            ("District:", assessment_data.get('district', 'N/A')),
            ("Facility Name:", assessment_data.get('facilityName', 'N/A')),
            ("Facility Level:", assessment_data.get('facilityLevel', 'N/A')),
            ("Ownership:", assessment_data.get('ownership', 'N/A')),
            ("Assessor Name:", assessment_data.get('assessorName', 'N/A')),
            ("Assessment Date:", assessment_data.get('assessmentDate', 'N/A'))
        ]
        
        for label, value in facility_info:
            ws.cell(row=row, column=1).value = label
            ws.cell(row=row, column=1).font = Font(bold=True)
            ws.cell(row=row, column=1).fill = info_fill
            ws.cell(row=row, column=1).border = thin_border
            
            ws.merge_cells(f'B{row}:C{row}')
            ws.cell(row=row, column=2).value = DataValidator.sanitize_input(value)
            ws.cell(row=row, column=2).border = thin_border
            ws.cell(row=row, column=3).border = thin_border
            row += 1
        
        # Add scoring summary
        row += 1
        ws.merge_cells(f'A{row}:C{row}')
        ws.cell(row=row, column=1).value = "ASSESSMENT SCORES"
        ws.cell(row=row, column=1).font = Font(size=14, bold=True, color="FFFFFF")
        ws.cell(row=row, column=1).fill = header_fill
        ws.cell(row=row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(row=row, column=1).border = thin_border
        ws.row_dimensions[row].height = 25
        row += 1
        
        if is_section_download:
            score_info = [
                ("Section Score:", f"{total_scored} / {total_possible}"),
                ("Percentage:", f"{percentage}%"),
                ("Performance Level:", performance)
            ]
        else:
            score_info = [
                ("Total Sections Assessed:", f"{scored_indicators} of {len(section_ids)}"),
                ("Maximum Possible Score:", f"{total_possible} points"),
                ("Total Score Achieved:", f"{total_scored} points"),
                ("Overall Percentage:", f"{percentage}%"),
                ("Performance Level:", performance)
            ]
        
        for label, value in score_info:
            ws.cell(row=row, column=1).value = label
            ws.cell(row=row, column=1).font = Font(bold=True)
            ws.cell(row=row, column=1).fill = info_fill
            ws.cell(row=row, column=1).border = thin_border
            
            ws.merge_cells(f'B{row}:C{row}')
            ws.cell(row=row, column=2).value = value
            ws.cell(row=row, column=2).font = Font(bold=True, size=12)
            if label == "Performance Level:":
                ws.cell(row=row, column=2).fill = PatternFill(start_color=perf_color, end_color=perf_color, fill_type="solid")
                ws.cell(row=row, column=2).font = Font(bold=True, size=12, color="FFFFFF" if perf_color in ["006400", "DC3545"] else "000000")
            ws.cell(row=row, column=2).border = thin_border
            ws.cell(row=row, column=2).alignment = Alignment(horizontal='center', vertical='center')
            ws.cell(row=row, column=3).border = thin_border
            row += 1
        
        # For section downloads, show questions with color-coded answers instead of legend
        if is_section_download:
            row += 1
            ws.merge_cells(f'A{row}:C{row}')
            ws.cell(row=row, column=1).value = "ASSESSMENT RESPONSES"
            ws.cell(row=row, column=1).font = Font(size=12, bold=True, color="FFFFFF")
            ws.cell(row=row, column=1).fill = header_fill
            ws.cell(row=row, column=1).alignment = Alignment(horizontal='center', vertical='center')
            ws.cell(row=row, column=1).border = thin_border
            row += 1
            
            # Get all scores/responses for this section
            section_key = sections_with_data[0]
            
            # Build question text map from section definitions FIRST
            question_text_map = {}
            section_definitions = assessment_data.get('section_definitions', {})
            
            if section_key in section_definitions:
                section_config = section_definitions[section_key]
                section_type = section_config.get('type', '')
                
                # Handle different section types
                # Type 1: Quality Matrix (services array) - NEW!
                if section_type == 'quality_matrix' and 'services' in section_config:
                    for service in section_config['services']:
                        question_text_map[service['id']] = service.get('name', service['id'])
                
                # Type 2: Register Checklist (registers array) - NEW!
                elif section_type == 'register_checklist' and 'registers' in section_config:
                    for register in section_config['registers']:
                        reg_id = register['id']
                        reg_name = register['name']
                        for column in section_config.get('columns', ['Available', 'Standard versions', '90% complete']):
                            indicator_id = f"{reg_id}_{column.lower().replace('%', '_pct').replace(' ', '_')}"
                            question_text_map[indicator_id] = f"{reg_name} - {column}"
                
                # Type 3: Questions array (most common)
                if 'questions' in section_config:
                    for q in section_config['questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                        
                        # Handle items within questions (multi_yes_no)
                        if 'items' in q and 'type' in q and q['type'] == 'multi_yes_no':
                            for idx, item in enumerate(q['items']):
                                # Map like tet_q1_hiv, tet_q1_syphilis
                                item_id = f"{q['id']}_{item.lower().replace(' ', '_')}"
                                question_text_map[item_id] = f"{q.get('text', '')} - {item}"
                        
                        # Handle fields within questions (data_entry_table)
                        if 'fields' in q:
                            for field in q['fields']:
                                # Use the field's name as question text
                                field_id = f"{section_key}_{field['id']}" if not field['id'].startswith(section_key) else field['id']
                                question_text_map[field_id] = field.get('name', field.get('label', field['id']))
                                # Also try without prefix
                                question_text_map[field['id']] = field.get('name', field.get('label', field['id']))
                        
                        # Handle sub-questions
                        if 'sub_questions' in q:
                            for sub_q in q['sub_questions']:
                                question_text_map[sub_q['id']] = sub_q.get('text', sub_q['id'])
                        
                        # Handle options (for checklist questions)
                        if 'options' in q:
                            for opt in q['options']:
                                if 'value' in opt:
                                    question_text_map[f"{q['id']}_{opt['value']}"] = f"{q.get('text', '')} - {opt.get('label', opt['value'])}"
                
                # Type 4: Inventory questions (supply chain)
                if 'inventory_questions' in section_config:
                    for q in section_config['inventory_questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                
                # Type 5: Scoring questions (supply chain)
                if 'scoring_questions' in section_config:
                    for q in section_config['scoring_questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                
                # Type 6: Old format indicators array - NEW!
                if 'indicators' in section_config:
                    for indicator in section_config['indicators']:
                        if isinstance(indicator, dict):
                            indicator_id = indicator.get('id', '')
                            indicator_text = indicator.get('text', indicator.get('name', indicator_id))
                            question_text_map[indicator_id] = indicator_text
            
            # NOW get all question keys - for section downloads, include ALL scores that have mappings
            # The question_text_map was built from the section definition, so any key in scores
            # that has a mapping is part of this section
            all_question_keys = []
            for key in scores.keys():
                # Skip the section aggregate score itself
                if key == section_key:
                    continue
                # If we have a text mapping for this key, it belongs to this section
                if key in question_text_map:
                    all_question_keys.append(key)
            
            # Also include any mapped keys that might not be in scores yet (for display purposes)
            for key in question_text_map.keys():
                if key not in all_question_keys and key != section_key:
                    all_question_keys.append(key)
            
            print(f"\n{'='*60}", flush=True)
            print(f"DEBUG: Section key: {section_key}", flush=True)
            print(f"DEBUG: Question text map size: {len(question_text_map)}", flush=True)
            print(f"DEBUG: All question keys: {len(all_question_keys)}", flush=True)
            print(f"DEBUG: Sample question keys: {all_question_keys[:5] if len(all_question_keys) > 5 else all_question_keys}", flush=True)
            
            # Display each question with its color-coded answer
            questions_shown = 0
            print(f"DEBUG: Starting question loop... ", flush=True)
            for question_key in sorted(all_question_keys):
                # Skip the final section score
                if question_key == section_key:
                    print(f"DEBUG: Skipping section score: {section_key}", flush=True)
                    continue
                    
                response = scores.get(question_key, '')
                print(f"DEBUG: Processing {question_key} = '{response}' (type: {type(response)})", flush=True)
                
                # For section downloads, show ALL questions even if unanswered
                # This ensures comprehensive reports showing full section structure
                
                # Get question text from map or generate from key
                if question_key in question_text_map:
                    question_text = question_text_map[question_key]
                    print(f"DEBUG: Found mapping for {question_key}: {question_text[:50]}...", flush=True)
                else:
                    # Fallback: Format the question key nicely
                    clean_key = question_key.replace(section_key + '_', '').replace('_', ' ').title()
                    question_text = f"{clean_key}"
                    print(f"DEBUG: No mapping for {question_key}, using: {question_text}", flush=True)
                
                print(f"DEBUG: Writing question to Excel row {row}", flush=True)
                
                # Question text in column A (merged with B)
                ws.merge_cells(f'A{row}:B{row}')
                ws.cell(row=row, column=1).value = question_text
                ws.cell(row=row, column=1).font = Font(bold=True, size=10)
                ws.cell(row=row, column=1).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                ws.cell(row=row, column=1).border = thin_border
                ws.cell(row=row, column=2).border = thin_border
                
                # Response in column C with color coding
                # Handle unanswered questions
                if response in ['', None, 'N/A']:
                    display_response = "Not Answered"
                else:
                    display_response = str(response)
                
                ws.cell(row=row, column=3).value = display_response
                ws.cell(row=row, column=3).font = Font(bold=True, size=11, color="000000")
                ws.cell(row=row, column=3).alignment = Alignment(horizontal='center', vertical='center')
                ws.cell(row=row, column=3).border = Border(
                    left=Side(style='medium'), right=Side(style='medium'),
                    top=Side(style='medium'), bottom=Side(style='medium')
                )
                
                # Color code based on response
                if response in ['', None, 'N/A']:
                    # Gray for unanswered questions
                    ws.cell(row=row, column=3).fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
                    ws.cell(row=row, column=3).font = Font(bold=True, size=11, color="666666", italic=True)
                else:
                    response_lower = str(response).lower()
                    if response_lower in ['yes', 'y']:
                        ws.cell(row=row, column=3).fill = PatternFill(start_color="006400", end_color="006400", fill_type="solid")
                        ws.cell(row=row, column=3).font = Font(bold=True, size=11, color="FFFFFF")
                    elif response_lower in ['no', 'n']:
                        ws.cell(row=row, column=3).fill = PatternFill(start_color="DC3545", end_color="DC3545", fill_type="solid")
                        ws.cell(row=row, column=3).font = Font(bold=True, size=11, color="FFFFFF")
                    else:
                        # For numeric or percentage values, use light blue
                        ws.cell(row=row, column=3).fill = PatternFill(start_color="E8F4F8", end_color="E8F4F8", fill_type="solid")
                
                questions_shown += 1
                row += 1
            
            print(f"DEBUG: *** Total questions shown: {questions_shown} ***", flush=True)
            print(f"{'='*60}\n", flush=True)
        else:
            # Add legend for full reports
            row += 1
            ws.merge_cells(f'A{row}:C{row}')
            ws.cell(row=row, column=1).value = "SCORING LEGEND"
            ws.cell(row=row, column=1).font = Font(size=12, bold=True, color="FFFFFF")
            ws.cell(row=row, column=1).fill = header_fill
            ws.cell(row=row, column=1).alignment = Alignment(horizontal='center', vertical='center')
            ws.cell(row=row, column=1).border = thin_border
            row += 1
            
            legend = [
                ("Score 1 (Red):", "Critical - Immediate action required"),
                ("Score 2 (Yellow):", "Needs improvement - Action required"),
                ("Score 3 (Light Green):", "Good - Minor improvements needed"),
                ("Score 4 (Dark Green):", "Excellent - Meets all standards")
            ]
            
            for label, desc in legend:
                ws.cell(row=row, column=1).value = label
                ws.cell(row=row, column=1).font = Font(bold=True)
                if "Red" in label:
                    ws.cell(row=row, column=1).fill = PatternFill(start_color="DC3545", end_color="DC3545", fill_type="solid")
                    ws.cell(row=row, column=1).font = Font(bold=True, color="FFFFFF")
                elif "Yellow" in label:
                    ws.cell(row=row, column=1).fill = PatternFill(start_color="FFC107", end_color="FFC107", fill_type="solid")
                elif "Light Green" in label:
                    ws.cell(row=row, column=1).fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
                elif "Dark Green" in label:
                    ws.cell(row=row, column=1).fill = PatternFill(start_color="006400", end_color="006400", fill_type="solid")
                    ws.cell(row=row, column=1).font = Font(bold=True, color="FFFFFF")
                ws.cell(row=row, column=1).border = thin_border
                
                ws.merge_cells(f'B{row}:C{row}')
                ws.cell(row=row, column=2).value = desc
                ws.cell(row=row, column=2).border = thin_border
                ws.cell(row=row, column=3).border = thin_border
                row += 1
    
    @staticmethod
    def _create_details_sheet(wb, assessment_data, is_section_download=False):
        """Create Assessment Details sheet with exact section-specific indicators"""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        ws = wb.create_sheet("Assessment Details")
        
        # Styling
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                            top=Side(style='thin'), bottom=Side(style='thin'))
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        section_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        
        # Get data
        scores = assessment_data.get('scores', {})
        comments = assessment_data.get('comments', {})
        section_definitions = assessment_data.get('section_definitions', {})
        
        # Identify sections with data
        sections_with_data = []
        for key in section_definitions.keys():
            if key in scores or any(k.startswith(key + '_') for k in scores.keys()):
                sections_with_data.append(key)
        
        # Column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20  
        ws.column_dimensions['C'].width = 50
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 40
        
        # Title
        ws.merge_cells('A1:E1')
        title_cell = ws['A1']
        if is_section_download and len(sections_with_data) == 1:
            section_config = section_definitions[sections_with_data[0]]
            title_cell.value = f"ASSESSMENT DETAILS: {section_config.get('name', sections_with_data[0])}"
        else:
            title_cell.value = "FACILITY ASSESSMENT DETAILED REPORT"
        title_cell.font = Font(size=16, bold=True, color="FFFFFF")
        title_cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        title_cell.border = thin_border
        
        # Facility info
        row = 3
        ws[f'A{row}'] = "District:"
        ws.merge_cells(f'B{row}:E{row}')
        ws[f'B{row}'] = assessment_data.get('district', 'N/A')
        
        row += 1
        ws[f'A{row}'] = "Facility Name:"
        ws.merge_cells(f'B{row}:E{row}')
        ws[f'B{row}'] = assessment_data.get('facilityName', 'N/A')
        
        row += 1
        ws[f'A{row}'] = "Assessor:"
        ws[f'B{row}'] = assessment_data.get('assessorName', 'N/A')
        ws[f'D{row}'] = "Date:"
        ws[f'E{row}'] = assessment_data.get('assessmentDate', 'N/A')
        
        # Headers
        row += 2
        headers = ['Section', 'Indicator ID', 'Indicator Name', 'Score', 'Comments']
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Process sections
        row += 1
        sections_to_process = [sections_with_data[0]] if is_section_download and len(sections_with_data) == 1 else sections_with_data
        
        for section_key in sections_to_process:
            if section_key not in section_definitions:
                continue
                
            section_config = section_definitions[section_key]
            section_name = section_config.get('name', section_key)
            
            # Section header
            ws.merge_cells(f'A{row}:E{row}')
            cell = ws.cell(row=row, column=1)
            cell.value = section_name
            cell.fill = section_fill
            cell.font = Font(bold=True, size=12)
            cell.border = thin_border
            row += 1
            
            # Extract indicators based on section type
            indicators_to_show = []
            
            # Type 1: register_checklist (ANC/Maternity/PNC Registers)
            if section_config.get('type') == 'register_checklist' and 'registers' in section_config:
                for register in section_config['registers']:
                    reg_id = register['id']
                    reg_name = register['name']
                    for column in section_config.get('columns', ['Available', 'Standard versions', '90% complete']):
                        indicator_id = f"{reg_id}_{column.lower().replace('%', '_pct').replace(' ', '_')}"
                        indicator_name = f"{reg_name} - {column}"
                        indicators_to_show.append((indicator_id, indicator_name))
            
            # Type 2: questions array
            elif 'questions' in section_config:
                for q in section_config['questions']:
                    q_id = q['id']
                    q_text = q.get('text', q_id)
                    
                    # multi_yes_no with items
                    if 'items' in q and q.get('type') == 'multi_yes_no':
                        for item in q['items']:
                            item_id = f"{q_id}_{item.lower().replace(' ', '_').replace('/', '_').replace('.', '')}"
                            indicators_to_show.append((item_id, f"{q_text} - {item}"))
                    
                    # data_entry_table with fields
                    elif 'fields' in q:
                        for field in q['fields']:
                            field_id = field['id']
                            field_name = field.get('name', field.get('label', field_id))
                            indicators_to_show.append((field_id, field_name))
                    
                    # checklist with items
                    elif q.get('type') == 'checklist' and 'items' in q:
                        for item in q['items']:
                            if isinstance(item, dict) and 'id' in item:
                                indicators_to_show.append((item['id'], item.get('text', item['id'])))
                    
                    # Regular question
                    else:
                        indicators_to_show.append((q_id, q_text))
            
            # Type 3: inventory_questions (supply chain)
            if 'inventory_questions' in section_config:
                for q in section_config['inventory_questions']:
                    indicators_to_show.append((q['id'], q.get('text', q['id'])))
            
            # Type 4: scoring_questions (supply chain)
            if 'scoring_questions' in section_config:
                for q in section_config['scoring_questions']:
                    indicators_to_show.append((q['id'], q.get('text', q['id'])))
            
            # Type 5: services (quality_matrix)
            if 'services' in section_config:
                for service in section_config['services']:
                    indicators_to_show.append((service['id'], service.get('name', service['id'])))
            
            # Display each indicator
            for indicator_id, indicator_name in indicators_to_show:
                score = scores.get(indicator_id, '')
                comment = comments.get(indicator_id, '')
                
                ws.cell(row=row, column=1).value = section_name
                ws.cell(row=row, column=2).value = indicator_id
                ws.cell(row=row, column=3).value = indicator_name
                ws.cell(row=row, column=4).value = score if score else ''
                ws.cell(row=row, column=5).value = comment
                
                # Apply borders
                for col_num in range(1, 6):
                    cell = ws.cell(row=row, column=col_num)
                    cell.border = thin_border
                    cell.alignment = Alignment(vertical='center', wrap_text=True)
                
                # Color code score
                score_cell = ws.cell(row=row, column=4)
                score_cell.alignment = Alignment(horizontal='center', vertical='center')
                if score and score not in ['', 'N/A', None]:
                    try:
                        score_val = int(float(score))
                        if score_val == 1:
                            score_cell.fill = PatternFill(start_color="DC3545", end_color="DC3545", fill_type="solid")
                            score_cell.font = Font(bold=True, color="FFFFFF")
                        elif score_val == 2:
                            score_cell.fill = PatternFill(start_color="FFC107", end_color="FFC107", fill_type="solid")
                            score_cell.font = Font(bold=True)
                        elif score_val == 3:
                            score_cell.fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
                            score_cell.font = Font(bold=True)
                        elif score_val == 4:
                            score_cell.fill = PatternFill(start_color="006400", end_color="006400", fill_type="solid")
                            score_cell.font = Font(bold=True, color="FFFFFF")
                    except:
                        pass
                
                row += 1
            
            # Section score
            section_score = scores.get(section_key, '')
            if section_score and section_score not in ['', 'N/A', None]:
                ws.cell(row=row, column=1).value = section_name
                ws.cell(row=row, column=2).value = section_key
                ws.cell(row=row, column=3).value = "SECTION SCORE"
                ws.cell(row=row, column=3).font = Font(bold=True, size=11)
                ws.cell(row=row, column=4).value = section_score
                ws.cell(row=row, column=5).value = ""
                
                for col_num in range(1, 6):
                    ws.cell(row=row, column=col_num).border = thin_border
                
                score_cell = ws.cell(row=row, column=4)
                score_cell.alignment = Alignment(horizontal='center', vertical='center')
                try:
                    score_val = int(float(section_score))
                    if score_val == 1:
                        score_cell.fill = PatternFill(start_color="DC3545", end_color="DC3545", fill_type="solid")
                        score_cell.font = Font(bold=True, size=11, color="FFFFFF")
                    elif score_val == 2:
                        score_cell.fill = PatternFill(start_color="FFC107", end_color="FFC107", fill_type="solid")
                        score_cell.font = Font(bold=True, size=11)
                    elif score_val == 3:
                        score_cell.fill = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
                        score_cell.font = Font(bold=True, size=11)
                    elif score_val == 4:
                        score_cell.fill = PatternFill(start_color="006400", end_color="006400", fill_type="solid")
                        score_cell.font = Font(bold=True, size=11, color="FFFFFF")
                except:
                    pass
                
                row += 1
            
            row += 1  # Space between sections


    def _create_action_plan_sheet(wb: Workbook, assessment_data: Dict):
        """Create action plan sheet with recommendations based on scores"""
        ws = wb.create_sheet("Action Plan")
        
        # Define styling
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        
        # Score-based color coding (matching assessment detail colors)
        score_1_fill = PatternFill(start_color="DC143C", end_color="DC143C", fill_type="solid")  # Dark Red - Critical
        score_2_fill = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")  # Red - Urgent
        score_3_fill = PatternFill(start_color="FFD93D", end_color="FFD93D", fill_type="solid")  # Yellow - Moderate
        score_4_fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")  # Orange - Minor
        
        # Set column widths
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 60
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 30
        
        # Add title
        ws.merge_cells('A1:D1')
        title_cell = ws['A1']
        title_cell.value = "IMPROVEMENT ACTION PLAN - Priority Areas for Intervention"
        title_cell.font = Font(size=16, bold=True, color="FFFFFF")
        title_cell.fill = header_fill
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        title_cell.border = thin_border
        
        # Add headers
        row = 3
        headers = ['Indicator Name', 'Recommended Action', 'Priority Level', 'Timeline']
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = Font(bold=True, size=11, color="FFFFFF")
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        # Define all sections and indicators (matching assessment details)
        sections = {
            "Service Delivery": [
                {"id": "sd1", "name": "Are hepatitis B testing services available and integrated in ANC, maternity, and postnatal care?", "max_score": 5},
                {"id": "sd2", "name": "Are HBV-positive pregnant women immediately enrolled for treatment/prophylaxis (TDF or TDF/3TC) at the facility?", "max_score": 5},
                {"id": "sd3", "name": "Is the hepatitis B birth dose vaccine (HepB-BD) routinely administered within 24 hours of birth?", "max_score": 5},
                {"id": "sd4", "name": "Are there documented standard operating procedures (SOPs) for PMTCT of hepatitis B at the facility?", "max_score": 5},
                {"id": "sd5", "name": "Does the facility have integrated maternity and ANC registers that capture hepatitis B screening, treatment, and HepB-BD vaccination?", "max_score": 5},
                {"id": "sd6", "name": "Are follow-up mechanisms in place for HBV-positive mothers and their exposed infants (post-vaccination serology at 9-12 months)?", "max_score": 5}
            ],
            "Human Resource and Service Delivery Points": [
                {"id": "hr1", "name": "What is the primary funding source for personnel delivering PMTCT services at this facility?", "max_score": 5},
                {"id": "hr2", "name": "If partner-supported, which partner(s) support PMTCT personnel?", "max_score": 5},
                {"id": "hr3", "name": "Are PMTCT services integrated and provided at a designated Mother-Baby Care Point (MBCP)?", "max_score": 5},
                {"id": "hr4", "name": "Where is hepatitis B testing for pregnant women primarily conducted?", "max_score": 5},
                {"id": "hr5", "name": "Where is treatment/prophylaxis provided for HBV", "max_score": 5},
                {"id": "hr6", "name": "Where is the Hepatitis B birth dose vaccine administered?", "max_score": 5},
                {"id": "hr7", "name": "Are there designated personnel responsible for providing comprehensive PMTCT services at each service delivery point?", "max_score": 5},
                {"id": "hr8", "name": "If PMTCT services are not co-located, what are the key gaps and how are patients referred between service points?", "max_score": 5}
            ],
            "Supply Chain Reliability for HBV commodities": [
                {"id": "sc1", "name": "Are HBsAg test kits currently in stock?", "max_score": 5},
                {"id": "sc2", "name": "How many months of stock are available for HBsAg test kits?", "max_score": 5},
                {"id": "sc3", "name": "Is TDF or TDF/3TC for prophylaxis currently in stock?", "max_score": 5},
                {"id": "sc4", "name": "How many months of stock are available for TDF or TDF/3TC?", "max_score": 5},
                {"id": "sc5", "name": "Are hepatitis B birth dose (HepB-BD) vaccines currently available in maternity?", "max_score": 5},
                {"id": "sc6", "name": "How many months of stock are available for HepB-BD vaccines?", "max_score": 5},
                {"id": "sc7", "name": "Has a stock-out of HBsAg test kits in the past 3 months caused missed screening of pregnant women?", "max_score": 5},
                {"id": "sc8", "name": "Has there been a stock-out of TDF or TDF/3TC for HBV-positive women in the last 3 months?", "max_score": 5},
                {"id": "sc9", "name": "Are hepatitis B birth dose vaccines available 24/7 in maternity units for timely newborn immunization?", "max_score": 5}
            ],
            "Infrastructure & Equipment": [
                {"id": "if1", "name": "Is there a functional cold chain (refrigerator/freezer) for storing hepatitis B vaccines at the correct temperature (2-8°C)?", "max_score": 5},
                {"id": "if2", "name": "Are rapid diagnostic test kits (RDTs) or laboratory equipment for HBsAg testing available and functional?", "max_score": 5},
                {"id": "if3", "name": "Does the facility have a designated, private space for counseling HBV-positive pregnant women?", "max_score": 5},
                {"id": "if4", "name": "Is there reliable power supply or backup generator to maintain cold chain and laboratory services?", "max_score": 5},
                {"id": "if5", "name": "Are there adequate supplies of safety boxes, gloves, and sharps disposal for safe HBV testing and vaccination?", "max_score": 5}
            ]
        }
        
        # Get scores and comments
        scores = assessment_data.get('scores', {})
        comments = assessment_data.get('comments', {})
        
        # Create indicator lookup dictionary
        indicator_lookup = {}
        for section_name, indicators in sections.items():
            for indicator in indicators:
                indicator_lookup[indicator['id']] = indicator['name']
        
        # Generate action plan for indicators with scores 1-4
        row += 1
        action_items = []
        
        for indicator_id, score in scores.items():
            if isinstance(score, (int, str)) and str(score).isdigit():
                score_int = int(score)
                if score_int <= 4:  # Include scores 1-4 in action plan
                    # Determine priority, timeline, and color based on score
                    if score_int == 1:
                        priority = "CRITICAL"
                        timeline = "Immediate (Within 2 weeks)"
                        fill_color = score_1_fill
                        font_color = "FFFFFF"
                    elif score_int == 2:
                        priority = "URGENT"
                        timeline = "Short-term (1-3 months)"
                        fill_color = score_2_fill
                        font_color = "FFFFFF"
                    elif score_int == 3:
                        priority = "MODERATE"
                        timeline = "Medium-term (3-6 months)"
                        fill_color = score_3_fill
                        font_color = "000000"
                    else:  # score == 4
                        priority = "MINOR"
                        timeline = "Long-term (6-12 months)"
                        fill_color = score_4_fill
                        font_color = "000000"
                    
                    # Get indicator name
                    indicator_name = indicator_lookup.get(indicator_id, indicator_id.upper())
                    
                    # Generate action based on comment or generic
                    comment = comments.get(indicator_id, "")
                    action = f"Address identified gaps: {comment}" if comment else f"Develop and implement improvement plan to enhance this indicator"
                    
                    action_items.append({
                        'name': indicator_name,
                        'action': action,
                        'priority': priority,
                        'timeline': timeline,
                        'fill': fill_color,
                        'font_color': font_color,
                        'score': score_int
                    })
        
        # Sort by score (lowest first - most critical)
        action_items.sort(key=lambda x: x['score'])
        
        # Populate action plan
        for item in action_items:
            ws.cell(row=row, column=1).value = item['name']
            ws.cell(row=row, column=2).value = item['action']
            ws.cell(row=row, column=3).value = item['priority']
            ws.cell(row=row, column=4).value = item['timeline']
            
            # Apply styling
            for col_num in range(1, 5):
                cell = ws.cell(row=row, column=col_num)
                cell.border = thin_border
                cell.alignment = Alignment(vertical='center', wrap_text=True)
                
                # Color code the priority column
                if col_num == 3:
                    cell.fill = item['fill']
                    cell.font = Font(bold=True, color=item['font_color'])
            
            row += 1
        
        # Add summary note if no action items
        if not action_items:
            ws.cell(row=row, column=1).value = "No critical action items identified. All indicators scored 5 (excellent performance)."
            ws.merge_cells(f'A{row}:D{row}')
            cell = ws.cell(row=row, column=1)
            cell.fill = PatternFill(start_color="6BCF7F", end_color="6BCF7F", fill_type="solid")
            cell.font = Font(bold=True, color="FFFFFF")
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    @staticmethod
    def _save_workbook(wb: Workbook, filename: str) -> str:
        """Save workbook to temporary file"""
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        wb.save(filepath)
        return filepath

class EmailService:
    """Handle email operations with proper error handling"""
    
    @staticmethod
    @log_function_call
    def send_email(filepath: str, filename: str, email_type: str = "registration") -> Tuple[bool, str]:
        """Send email with attachment"""
        try:
            # Validate configuration
            if not Config.SMTP_USERNAME or not Config.SMTP_PASSWORD:
                return False, "Email configuration not properly set"
            
            # Create message
            msg = MIMEMultipart()
            msg['From'] = Config.SMTP_USERNAME
            msg['To'] = Config.RECIPIENT_EMAIL
            
            # Set subject and body based on type
            if email_type == "assessment":
                msg['Subject'] = f'Facility Assessment Report - {datetime.now().strftime("%Y-%m-%d %H:%M")}'
                body = EmailService._get_assessment_email_body()
            else:
                msg['Subject'] = f'Participant Registration - {datetime.now().strftime("%Y-%m-%d %H:%M")}'
                body = EmailService._get_registration_email_body()
            
            msg.attach(MIMEText(body, 'plain'))
            
            # Attach file
            EmailService._attach_file(msg, filepath, filename)
            
            # Send email
            success = EmailService._send_smtp_message(msg)
            
            if success:
                log_email_event("Email sent", Config.RECIPIENT_EMAIL, True, f"File: {filename}")
                return True, "Email sent successfully"
            else:
                log_email_event("Email failed", Config.RECIPIENT_EMAIL, False, "SMTP error")
                return False, "Failed to send email"
                
        except Exception as e:
            log_email_event("Email error", Config.RECIPIENT_EMAIL, False, str(e))
            return False, f"Email error: {str(e)}"
    
    @staticmethod
    def _get_registration_email_body() -> str:
        """Get registration email body"""
        return f"""
Dear Team,

Please find attached the participant registration data.

Registration Date: {datetime.now().strftime("%Y-%m-%d at %H:%M")}

Best regards,
Registration System
"""
    
    @staticmethod
    def _get_assessment_email_body() -> str:
        """Get assessment email body"""
        return f"""
Dear Team,

Please find attached the facility assessment report.

Assessment Date: {datetime.now().strftime("%Y-%m-%d at %H:%M")}

This report includes:
- Overall facility performance summary
- Detailed indicator scores
- Recommended action plan

Best regards,
Assessment System
"""
    
    @staticmethod
    def _attach_file(msg: MIMEMultipart, filepath: str, filename: str):
        """Attach file to email message"""
        with open(filepath, 'rb') as f:
            part = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={filename}')
            msg.attach(part)
    
    @staticmethod
    def _send_smtp_message(msg: MIMEMultipart) -> bool:
        """Send message via SMTP"""
        try:
            server = smtplib.SMTP(Config.SMTP_SERVER, Config.SMTP_PORT)
            if Config.SMTP_USE_TLS:
                server.starttls()
            server.login(Config.SMTP_USERNAME, Config.SMTP_PASSWORD)
            server.send_message(msg)
            server.quit()
            return True
        except Exception as e:
            print(f"SMTP Error: {e}")
            return False

class FileManager:
    """Handle file operations safely"""
    
    @staticmethod
    @log_function_call
    def cleanup_temp_file(filepath: str) -> bool:
        """Safely remove temporary file"""
        try:
            if filepath and os.path.exists(filepath):
                os.remove(filepath)
                return True
        except Exception as e:
            print(f"Warning: Could not delete temp file {filepath}: {e}")
        return False
    
    @staticmethod
    @log_function_call
    def ensure_temp_directory() -> str:
        """Ensure temporary directory exists"""
        temp_dir = tempfile.gettempdir()
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        return temp_dir


class WordGenerator:
    """Handle Word document generation with proper error handling"""
    
    @staticmethod
    @log_function_call
    def create_assessment_word(assessment_data: Dict) -> Tuple[str, str]:
        """Generate comprehensive Word assessment report
        
        Args:
            assessment_data: Dictionary containing assessment information
        """
        try:
            doc = Document()
            
            # Detect if this is a section-specific download
            scores = assessment_data.get('scores', {})
            section_definitions = assessment_data.get('section_definitions', {})
            # Use ALL section keys from section_definitions instead of hardcoded list
            section_ids = list(section_definitions.keys())
            sections_with_data = [key for key in section_ids if key in scores or any(k.startswith(key + '_') for k in scores.keys())]
            is_section_download = len(sections_with_data) == 1
            
            # Create summary section
            WordGenerator._create_summary_section(doc, assessment_data, is_section_download)
            
            # Add page break
            doc.add_page_break()
            
            # Create detailed scores section
            WordGenerator._create_details_section(doc, assessment_data, is_section_download)
            
            # Add page break
            doc.add_page_break()
            
            # Create action plan section
            WordGenerator._create_action_plan_section(doc, assessment_data)
            
            # Generate filename and save
            facility_name = assessment_data.get('facilityName', 'Unknown').replace(' ', '_')
            filename = f'assessment_report_{facility_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            filepath = WordGenerator._save_document(doc, filename)
            
            return filepath, filename
            
        except Exception as e:
            raise Exception(f"Failed to create assessment Word file: {str(e)}")
    
    @staticmethod
    def _create_summary_section(doc: Document, assessment_data: Dict, is_section_download: bool = False):
        """Create comprehensive summary section"""
        scores = assessment_data.get('scores', {})
        section_ids = ['triple_elimination_treatment', 'art_pmtct', 'quality_pmtct', 'patient_tracking',
                      'adherence_support', 'facility_linkage', 'sti_screening', 'early_infant_diagnosis',
                      'ctx_hei', 'tracking_hei', 'enrolment_eid_art', 'hei_eid_registers',
                      'supply_chain_eid', 'supply_chain_pmtct', 'supply_chain_syphilis', 'supply_chain_hepb',
                      'patient_records']
        
        sections_with_data = [key for key in section_ids if key in scores or any(k.startswith(key + '_') for k in scores.keys())]
        
        # Section name mapping
        section_name_map = {
            'triple_elimination_treatment': 'TRIPLE ELIMINATION (HIV, Syphilis and Hep B): Linkage to treatment',
            'art_pmtct': 'ART in PMTCT Facilities (VL uptake at ANC1 for ARTK)',
            'quality_pmtct': 'Quality of PMTCT services (Review 10 client\'s charts)',
            'patient_tracking': 'Patient Tracking HIV+ Pregnant Women',
            'adherence_support': 'Adherence Support',
            'facility_linkage': 'Facility Linkage to Community Care and Support Services for Adult PLHIV',
            'sti_screening': 'STI Screening and Management in HIV Clinics Serving General Population',
            'early_infant_diagnosis': 'Early Infant Diagnosis [HEI]',
            'ctx_hei': 'CTX for HIV-Exposed Infants [HEI]',
            'tracking_hei': 'Tracking HIV-Exposed Infants [HEI]',
            'enrolment_eid_art': 'Enrolment of HIV-Infected Infants (EID Services into ART Services)',
            'hei_eid_registers': 'HIV Exposed Infant/Early Infant Diagnosis Registers',
            'supply_chain_eid': 'Supply Chain Reliability (Early Infant Diagnosis) [HEI]',
            'supply_chain_pmtct': 'Supply Chain Reliability – HIV PMTCT (Mother)',
            'supply_chain_syphilis': 'Supply Chain Reliability – Syphilis PMTCT',
            'supply_chain_hepb': 'Supply Chain Reliability – Hepatitis B PMTCT (Maternal and Birth Dose)',
            'patient_records': 'Patient/Beneficiary Records'
        }
        
        # Calculate scores
        total_possible = 0
        total_scored = 0
        scored_indicators = 0
        section_name = None
        
        if is_section_download:
            section_key = sections_with_data[0]
            section_name = section_name_map.get(section_key, section_key.upper())
            total_possible = 4
            score = scores.get(section_key, 0)
            try:
                total_scored = float(score) if score not in ['N/A', '', None] else 0
            except (ValueError, TypeError):
                total_scored = 0
            scored_indicators = 1 if total_scored > 0 else 0
        else:
            for section_id in section_ids:
                total_possible += 4
                score = scores.get(section_id, 0)
                try:
                    score_val = float(score) if score not in ['N/A', '', None] else 0
                    if score_val > 0:
                        total_scored += score_val
                        scored_indicators += 1
                except (ValueError, TypeError):
                    pass
        
        percentage = round((total_scored / total_possible * 100), 1) if total_possible > 0 else 0
        
        # Determine performance level and color
        if percentage >= 75:
            performance = "EXCELLENT"
            perf_color = RGBColor(0, 100, 0)  # Dark Green
        elif percentage >= 50:
            performance = "GOOD"
            perf_color = RGBColor(144, 238, 144)  # Light Green
        elif percentage >= 25:
            performance = "NEEDS IMPROVEMENT"
            perf_color = RGBColor(255, 193, 7)  # Yellow
        else:
            performance = "CRITICAL"
            perf_color = RGBColor(220, 53, 69)  # Red
        
        # Add title
        title = doc.add_heading('', level=0)
        title_run = title.add_run('ASSESSMENT SUMMARY\n' if not is_section_download else f'ASSESSMENT SUMMARY\n{section_name}')
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add facility information
        doc.add_paragraph()
        facility_info = [
            ("District:", assessment_data.get('district', 'N/A')),
            ("Facility Name:", assessment_data.get('facilityName', 'N/A')),
            ("Facility Level:", assessment_data.get('facilityLevel', 'N/A')),
            ("Ownership:", assessment_data.get('ownership', 'N/A')),
            ("Assessor Name:", assessment_data.get('assessorName', 'N/A')),
            ("Assessment Date:", assessment_data.get('assessmentDate', 'N/A'))
        ]
        
        for label, value in facility_info:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            p.add_run(f' {DataValidator.sanitize_input(value)}')
        
        # Add scoring summary
        doc.add_paragraph()
        score_heading = doc.add_heading('ASSESSMENT SCORES', level=2)
        score_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if is_section_download:
            score_info = [
                ("Section Score:", f"{total_scored} / {total_possible}"),
                ("Percentage:", f"{percentage}%"),
                ("Performance Level:", performance)
            ]
        else:
            score_info = [
                ("Total Sections Assessed:", f"{scored_indicators} of {len(section_ids)}"),
                ("Maximum Possible Score:", f"{total_possible} points"),
                ("Total Score Achieved:", f"{total_scored} points"),
                ("Overall Percentage:", f"{percentage}%"),
                ("Performance Level:", performance)
            ]
        
        for label, value in score_info:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            run = p.add_run(f' {value}')
            run.font.size = Pt(12)
            run.bold = True
            if label == "Performance Level:":
                run.font.color.rgb = perf_color
        
        # Show questions with color-coded answers for all downloads
        doc.add_paragraph()
        responses_heading = doc.add_heading('ASSESSMENT RESPONSES', level=2)
        responses_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Build question text map from section definitions FIRST
        question_text_map = {}
        section_definitions = assessment_data.get('section_definitions', {})
        
        # Determine which sections to process
        if is_section_download:
            section_key = sections_with_data[0]
            sections_to_process = [section_key]
        else:
            # For full reports, process all sections
            sections_to_process = section_definitions.keys()
        
        for current_section_key in sections_to_process:
            if current_section_key in section_definitions:
                section_config = section_definitions[current_section_key]
                section_type = section_config.get('type', '')
                
                # Type 1: Quality Matrix (services array) - NEW!
                if section_type == 'quality_matrix' and 'services' in section_config:
                    for service in section_config['services']:
                        question_text_map[service['id']] = service.get('name', service['id'])
                
                # Type 2: Register Checklist (registers array) - NEW!
                elif section_type == 'register_checklist' and 'registers' in section_config:
                    for register in section_config['registers']:
                        reg_id = register['id']
                        reg_name = register['name']
                        for column in section_config.get('columns', ['Available', 'Standard versions', '90% complete']):
                            indicator_id = f"{reg_id}_{column.lower().replace('%', '_pct').replace(' ', '_')}"
                            question_text_map[indicator_id] = f"{reg_name} - {column}"
                
                # Type 3: Questions array (most common)
                if 'questions' in section_config:
                    for q in section_config['questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                        
                        if 'items' in q and 'type' in q and q['type'] == 'multi_yes_no':
                            for item in q['items']:
                                item_id = f"{q['id']}_{item.lower().replace(' ', '_')}"
                                question_text_map[item_id] = f"{q.get('text', '')} - {item}"
                        
                        if 'fields' in q:
                            for field in q['fields']:
                                field_id = f"{current_section_key}_{field['id']}" if not field['id'].startswith(current_section_key) else field['id']
                                question_text_map[field_id] = field.get('name', field.get('label', field['id']))
                                question_text_map[field['id']] = field.get('name', field.get('label', field['id']))
                        
                        if 'sub_questions' in q:
                            for sub_q in q['sub_questions']:
                                question_text_map[sub_q['id']] = sub_q.get('text', sub_q['id'])
                        
                        if 'options' in q:
                            for opt in q['options']:
                                if 'value' in opt:
                                    question_text_map[f"{q['id']}_{opt['value']}"] = f"{q.get('text', '')} - {opt.get('label', opt['value'])}"
                
                # Type 4: Inventory questions (supply chain)
                if 'inventory_questions' in section_config:
                    for q in section_config['inventory_questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                
                # Type 5: Scoring questions (supply chain)
                if 'scoring_questions' in section_config:
                    for q in section_config['scoring_questions']:
                        question_text_map[q['id']] = q.get('text', q['id'])
                
                # Type 6: Old format indicators array - NEW!
                if 'indicators' in section_config:
                    for indicator in section_config['indicators']:
                        if isinstance(indicator, dict):
                            indicator_id = indicator.get('id', '')
                            indicator_text = indicator.get('text', indicator.get('name', indicator_id))
                            question_text_map[indicator_id] = indicator_text
        
        # NOW get all question keys based on download type
        if is_section_download:
            # For section downloads, include ALL scores that have mappings from the section definition
            # The question_text_map was built from the section definition, so any key in scores
            # that has a mapping is part of this section
            all_question_keys = []
            section_key = sections_with_data[0]
            for key in scores.keys():
                # Skip the section aggregate score itself
                if key == section_key:
                    continue
                # If we have a text mapping for this key, it belongs to this section
                if key in question_text_map:
                    all_question_keys.append(key)
            
            # Also include any mapped keys that might not be in scores yet (for completeness)
            for key in question_text_map.keys():
                if key not in all_question_keys and key != section_key:
                    all_question_keys.append(key)
        else:
            # For full reports, use all mapped keys and scored keys
            all_question_keys = list(question_text_map.keys())
            section_ids = ['triple_elimination_treatment', 'art_pmtct', 'quality_pmtct', 'patient_tracking',
                          'adherence_support', 'facility_linkage', 'sti_screening', 'early_infant_diagnosis',
                          'ctx_hei', 'tracking_hei', 'enrolment_eid_art', 'hei_eid_registers',
                          'supply_chain_eid', 'supply_chain_pmtct', 'supply_chain_syphilis', 'supply_chain_hepb',
                          'patient_records']
            for key in scores.keys():
                if key not in all_question_keys and key not in section_ids:
                    all_question_keys.append(key)
        
        # Create table for responses
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Response'
        for cell in hdr_cells:
            WordGenerator._set_cell_color(cell, 31, 78, 120)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # Display each question with its color-coded answer
        for question_key in sorted(all_question_keys):
            # Skip section-level scores (already filtered out above)
            if question_key in ['triple_elimination_treatment', 'art_pmtct', 'quality_pmtct', 'patient_tracking',
                               'adherence_support', 'facility_linkage', 'sti_screening', 'early_infant_diagnosis',
                               'ctx_hei', 'tracking_hei', 'enrolment_eid_art', 'hei_eid_registers',
                               'supply_chain_eid', 'supply_chain_pmtct', 'supply_chain_syphilis', 'supply_chain_hepb',
                               'patient_records']:
                continue
                
            response = scores.get(question_key, '')
            
            # For section downloads, show ALL questions even if unanswered
            # This ensures comprehensive reports showing full section structure
            
            # Get question text from map or generate from key
            if question_key in question_text_map:
                question_text = question_text_map[question_key]
            else:
                clean_key = question_key.replace('_', ' ').title()
                question_text = f"{clean_key}"
            
            # Add row
            row_cells = table.add_row().cells
            row_cells[0].text = question_text
            
            # Handle unanswered questions
            if response in ['', None, 'N/A']:
                row_cells[1].text = "Not Answered"
                # Gray for unanswered questions
                WordGenerator._set_cell_color(row_cells[1], 211, 211, 211)
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(102, 102, 102)
                        run.font.italic = True
                        run.font.bold = True
            else:
                row_cells[1].text = str(response)
                
                # Color code based on response
                response_lower = str(response).lower()
                if response_lower in ['yes', 'y']:
                    WordGenerator._set_cell_color(row_cells[1], 0, 100, 0)
                    for paragraph in row_cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                elif response_lower in ['no', 'n']:
                    WordGenerator._set_cell_color(row_cells[1], 220, 53, 69)
                    for paragraph in row_cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                else:
                    WordGenerator._set_cell_color(row_cells[1], 232, 244, 248)
    
    @staticmethod
    def _create_details_section(doc: Document, assessment_data: Dict, is_section_download: bool = False):
        """Create comprehensive detailed scores section"""
        scores = assessment_data.get('scores', {})
        comments = assessment_data.get('comments', {})
        
        section_keys = ['triple_elimination_treatment', 'art_pmtct', 'quality_pmtct', 'patient_tracking',
                       'adherence_support', 'facility_linkage', 'sti_screening', 'early_infant_diagnosis',
                       'ctx_hei', 'tracking_hei', 'enrolment_eid_art', 'hei_eid_registers',
                       'supply_chain_eid', 'supply_chain_pmtct', 'supply_chain_syphilis', 'supply_chain_hepb',
                       'patient_records']
        
        sections_with_data = [key for key in section_keys if key in scores or any(k.startswith(key + '_') for k in scores.keys())]
        
        section_name_map = {
            'triple_elimination_treatment': 'TRIPLE ELIMINATION (HIV, Syphilis and Hep B): Linkage to treatment',
            'art_pmtct': 'ART in PMTCT Facilities (VL uptake at ANC1 for ARTK)',
            'quality_pmtct': 'Quality of PMTCT services (Review 10 client\'s charts)',
            'patient_tracking': 'Patient Tracking HIV+ Pregnant Women',
            'adherence_support': 'Adherence Support',
            'facility_linkage': 'Facility Linkage to Community Care and Support Services for Adult PLHIV',
            'sti_screening': 'STI Screening and Management in HIV Clinics Serving General Population',
            'early_infant_diagnosis': 'Early Infant Diagnosis [HEI]',
            'ctx_hei': 'CTX for HIV-Exposed Infants [HEI]',
            'tracking_hei': 'Tracking HIV-Exposed Infants [HEI]',
            'enrolment_eid_art': 'Enrolment of HIV-Infected Infants (EID Services into ART Services)',
            'hei_eid_registers': 'HIV Exposed Infant/Early Infant Diagnosis Registers',
            'supply_chain_eid': 'Supply Chain Reliability (Early Infant Diagnosis) [HEI]',
            'supply_chain_pmtct': 'Supply Chain Reliability – HIV PMTCT (Mother)',
            'supply_chain_syphilis': 'Supply Chain Reliability – Syphilis PMTCT',
            'supply_chain_hepb': 'Supply Chain Reliability – Hepatitis B PMTCT (Maternal and Birth Dose)',
            'patient_records': 'Patient/Beneficiary Records'
        }
        
        # Add title
        title = doc.add_heading('ASSESSMENT DETAILS', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add facility information
        p = doc.add_paragraph()
        p.add_run(f"District: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('district', 'N/A')))
        
        p = doc.add_paragraph()
        p.add_run(f"Facility Name: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('facilityName', 'N/A')))
        
        p = doc.add_paragraph()
        p.add_run(f"Facility Level: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('facilityLevel', 'N/A')))
        p.add_run("  |  ")
        p.add_run("Ownership: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('ownership', 'N/A')))
        
        p = doc.add_paragraph()
        p.add_run(f"Assessor: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('assessorName', 'N/A')))
        p.add_run("  |  ")
        p.add_run("Date: ").bold = True
        p.add_run(DataValidator.sanitize_input(assessment_data.get('assessmentDate', 'N/A')))
        
        doc.add_paragraph()
        
        # Filter sections to show based on download type
        if is_section_download:
            section_key = sections_with_data[0]
            target_section_name = section_name_map.get(section_key, section_key.upper())
            sections_to_show = [(target_section_name, section_key)]
        else:
            sections_to_show = [(section_name_map.get(key, key.upper()), key) for key in section_keys]
        
        # Create table for all sections
        for section_name, section_key in sections_to_show:
            # Check if section has data
            section_score = scores.get(section_key)
            if section_score in [None, '', 'N/A'] and not is_section_download:
                continue
            
            # Add section heading
            section_heading = doc.add_heading(section_name, level=2)
            
            # Get all keys for this section
            all_section_keys = [k for k in scores.keys() if k == section_key or k.startswith(section_key + '_')]
            
            if all_section_keys:
                # Create table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Indicator'
                hdr_cells[1].text = 'Score'
                hdr_cells[2].text = 'Comments'
                
                for cell in hdr_cells:
                    WordGenerator._set_cell_color(cell, 68, 114, 196)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                
                # Add data rows
                for key in sorted(all_section_keys):
                    score = scores.get(key, 'N/A')
                    comment = DataValidator.sanitize_input(comments.get(key, ''))
                    
                    if score in ['N/A', '', None] and not comment:
                        continue
                    
                    indicator_name = key.replace('_', ' ').title()
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = indicator_name
                    row_cells[1].text = str(score) if score != 'N/A' else ''
                    row_cells[2].text = comment
                    
                    # Color code the score cell
                    if score != 'N/A' and isinstance(score, (int, str, float)):
                        try:
                            score_value = int(float(score))
                            if score_value == 1:
                                WordGenerator._set_cell_color(row_cells[1], 220, 53, 69)
                                for paragraph in row_cells[1].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        run.font.bold = True
                            elif score_value == 2:
                                WordGenerator._set_cell_color(row_cells[1], 255, 193, 7)
                                for paragraph in row_cells[1].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                            elif score_value == 3:
                                WordGenerator._set_cell_color(row_cells[1], 144, 238, 144)
                                for paragraph in row_cells[1].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                            elif score_value == 4:
                                WordGenerator._set_cell_color(row_cells[1], 0, 100, 0)
                                for paragraph in row_cells[1].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        run.font.bold = True
                        except (ValueError, TypeError):
                            pass
                
                doc.add_paragraph()
    
    @staticmethod
    def _create_action_plan_section(doc: Document, assessment_data: Dict):
        """Create action plan section with recommendations"""
        scores = assessment_data.get('scores', {})
        comments = assessment_data.get('comments', {})
        
        # Add title
        title = doc.add_heading('IMPROVEMENT ACTION PLAN', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Priority Areas for Intervention')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].italic = True
        
        doc.add_paragraph()
        
        # Generate action plan for indicators with scores 1-4
        action_items = []
        
        for indicator_id, score in scores.items():
            if isinstance(score, (int, str)) and str(score).isdigit():
                score_int = int(score)
                if score_int <= 4:
                    # Determine priority, timeline, and color based on score
                    if score_int == 1:
                        priority = "CRITICAL"
                        timeline = "Immediate (Within 2 weeks)"
                        color = (220, 20, 60)
                    elif score_int == 2:
                        priority = "URGENT"
                        timeline = "Short-term (1-3 months)"
                        color = (255, 107, 107)
                    elif score_int == 3:
                        priority = "MODERATE"
                        timeline = "Medium-term (3-6 months)"
                        color = (255, 217, 61)
                    else:  # score == 4
                        priority = "MINOR"
                        timeline = "Long-term (6-12 months)"
                        color = (255, 165, 0)
                    
                    indicator_name = indicator_id.replace('_', ' ').title()
                    comment = comments.get(indicator_id, "")
                    action = f"Address identified gaps: {comment}" if comment else f"Develop and implement improvement plan to enhance this indicator"
                    
                    action_items.append({
                        'name': indicator_name,
                        'action': action,
                        'priority': priority,
                        'timeline': timeline,
                        'color': color,
                        'score': score_int
                    })
        
        # Sort by score (lowest first - most critical)
        action_items.sort(key=lambda x: x['score'])
        
        if action_items:
            # Create table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Indicator Name'
            hdr_cells[1].text = 'Recommended Action'
            hdr_cells[2].text = 'Priority Level'
            hdr_cells[3].text = 'Timeline'
            
            for cell in hdr_cells:
                WordGenerator._set_cell_color(cell, 31, 78, 120)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            
            # Populate action plan
            for item in action_items:
                row_cells = table.add_row().cells
                row_cells[0].text = item['name']
                row_cells[1].text = item['action']
                row_cells[2].text = item['priority']
                row_cells[3].text = item['timeline']
                
                # Color code the priority column
                WordGenerator._set_cell_color(row_cells[2], *item['color'])
                for paragraph in row_cells[2].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255) if item['score'] <= 2 else RGBColor(0, 0, 0)
                        run.font.bold = True
        else:
            p = doc.add_paragraph("No critical action items identified. All indicators scored excellently!")
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    @staticmethod
    def _set_cell_color(cell, r: int, g: int, b: int):
        """Set cell background color using RGB values"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), f"{r:02x}{g:02x}{b:02x}")
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    @staticmethod
    def _save_document(doc: Document, filename: str) -> str:
        """Save document to temporary file"""
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        return filepath


class PDFGenerator:
    """Handle PDF generation - creates an exact visual snapshot of the assessment section"""
    
    @staticmethod
    @log_function_call
    def create_section_pdf(assessment_data: Dict) -> Tuple[str, str]:
        """Generate PDF snapshot of assessment section exactly as it appears on screen
        
        Args:
            assessment_data: Dictionary containing assessment information with section HTML
        """
        try:
            from xhtml2pdf import pisa
            
            # Get data
            facility_name = assessment_data.get('facilityName', 'Unknown')
            district = assessment_data.get('district', 'N/A')
            assessor_name = assessment_data.get('assessorName', 'N/A')
            assessment_date = assessment_data.get('assessmentDate', 'N/A')
            facility_level = assessment_data.get('facilityLevel', 'N/A')
            ownership = assessment_data.get('ownership', 'N/A')
            section_html = assessment_data.get('sectionHTML', '')
            
            # Get section info
            scores = assessment_data.get('scores', {})
            section_definitions = assessment_data.get('section_definitions', {})
            
            # Identify the section
            section_key = None
            for key in section_definitions.keys():
                if key in scores or any(k.startswith(key + '_') for k in scores.keys()):
                    section_key = key
                    break
            
            if not section_key:
                raise Exception("No section found in scores")
            
            section_config = section_definitions.get(section_key, {})
            section_name = section_config.get('name', 'Assessment Section')
            
            # Get scores and comments
            scores_dict = assessment_data.get('scores', {})
            comments_dict = assessment_data.get('comments', {})
            
            # Create PDF file
            temp_dir = tempfile.gettempdir()
            filename = f'assessment_section_{facility_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            filepath = os.path.join(temp_dir, filename)
            
            # Build HTML content that matches screen appearance
            html_content = PDFGenerator._build_section_html(
                section_name, section_config, section_key,
                district, facility_name, facility_level, ownership,
                assessor_name, assessment_date,
                scores_dict, comments_dict, section_html
            )
            
            # Generate PDF from HTML using xhtml2pdf (reliable)
            with open(filepath, 'wb') as pdf_file:
                pisa_status = pisa.CreatePDF(
                    html_content.encode('utf-8'),
                    dest=pdf_file,
                    encoding='utf-8'
                )
                
                if pisa_status.err:
                    raise Exception(f"PDF generation had {pisa_status.err} errors")
            
            # Verify file was created
            if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
                raise Exception("PDF file was not created or is empty")
            
            return filepath, filename
            
        except Exception as e:
            raise Exception(f"Failed to create PDF file: {str(e)}")
    
    @staticmethod
    def _build_section_html(section_name, section_config, section_key,
                           district, facility_name, facility_level, ownership,
                           assessor_name, assessment_date,
                           scores, comments, section_html_content=''):
        """Build HTML content that exactly matches the screen appearance"""
        
        # Build the section content dynamically based on section structure
        section_content_html = PDFGenerator._generate_section_content(
            section_config, section_key, scores, comments
        )
        
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: A4;
            margin: 1cm;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #333;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            margin-bottom: 20px;
            border-radius: 8px;
        }}
        .header h1 {{
            margin: 0 0 10px 0;
            font-size: 24pt;
        }}
        .header h2 {{
            margin: 0;
            font-size: 16pt;
            font-weight: normal;
            opacity: 0.95;
        }}
        .facility-info {{
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            padding: 15px;
            margin-bottom: 20px;
            border-radius: 5px;
        }}
        .facility-info table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .facility-info td {{
            padding: 5px 10px;
        }}
        .facility-info td:first-child {{
            font-weight: bold;
            width: 150px;
            color: #495057;
        }}
        .standard-box {{
            background: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 12px;
            margin: 15px 0;
            border-radius: 4px;
        }}
        .standard-box strong {{
            color: #856404;
        }}
        .scoring-box {{
            background: #fff3cd;
            border: 1px solid #ffc107;
            padding: 15px;
            margin: 15px 0;
            border-radius: 5px;
        }}
        .scoring-box h3 {{
            margin: 0 0 10px 0;
            color: #856404;
        }}
        .scoring-box ul {{
            margin: 5px 0;
            padding-left: 20px;
        }}
        .scoring-box li {{
            margin: 5px 0;
        }}
        .red-dot {{ color: #DC3545; font-weight: bold; }}
        .yellow-dot {{ color: #FFC107; font-weight: bold; }}
        .light-green-dot {{ color: #90EE90; font-weight: bold; }}
        .dark-green-dot {{ color: #006400; font-weight: bold; }}
        
        .question-section {{
            margin: 20px 0;
        }}
        .question {{
            background: white;
            border: 1px solid #dee2e6;
            border-left: 4px solid #4472C4;
            padding: 15px;
            margin: 10px 0;
            border-radius: 4px;
        }}
        .question-label {{
            font-weight: bold;
            color: #4472C4;
            margin-bottom: 8px;
        }}
        .radio-options {{
            display: flex;
            gap: 20px;
            margin: 8px 0;
        }}
        .radio-option {{
            padding: 8px 16px;
            border: 2px solid #dee2e6;
            border-radius: 4px;
            background: white;
        }}
        .radio-option.selected {{
            background: #4472C4;
            color: white;
            border-color: #4472C4;
            font-weight: bold;
        }}
        table.data-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        table.data-table th {{
            background: #4472C4;
            color: white;
            padding: 10px;
            text-align: left;
            font-weight: bold;
        }}
        table.data-table td {{
            border: 1px solid #dee2e6;
            padding: 8px;
        }}
        table.data-table tr:nth-child(even) {{
            background: #f8f9fa;
        }}
        .score-badge {{
            display: inline-block;
            padding: 8px 16px;
            border-radius: 20px;
            font-weight: bold;
            margin: 10px 0;
        }}
        .score-red {{ background: #DC3545; color: white; }}
        .score-yellow {{ background: #FFC107; color: black; }}
        .score-light-green {{ background: #90EE90; color: black; }}
        .score-dark-green {{ background: #006400; color: white; }}
        
        .section-score {{
            margin-top: 20px;
            padding: 15px;
            background: #e7f3ff;
            border-radius: 5px;
            text-align: center;
        }}
        .section-score h3 {{
            margin: 0 0 5px 0;
            color: #1F4E78;
        }}
        .comments-box {{
            margin-top: 10px;
            padding: 10px;
            background: #f8f9fa;
            border-radius: 4px;
            font-style: italic;
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{section_name}</h1>
        <h2>Facility Assessment Report</h2>
    </div>
    
    <div class="facility-info">
        <table>
            <tr>
                <td>District:</td>
                <td>{district}</td>
                <td>Facility Level:</td>
                <td>{facility_level}</td>
            </tr>
            <tr>
                <td>Facility Name:</td>
                <td>{facility_name}</td>
                <td>Ownership:</td>
                <td>{ownership}</td>
            </tr>
            <tr>
                <td>Assessor:</td>
                <td>{assessor_name}</td>
                <td>Assessment Date:</td>
                <td>{assessment_date}</td>
            </tr>
        </table>
    </div>
    
    {section_content_html}
</body>
</html>
        """
        return html
    
    @staticmethod
    def _generate_section_content(section_config, section_key, scores, comments):
        """Generate HTML content for the section based on its type and structure"""
        content = ""
        section_type = section_config.get('type', '')
        
        # Add standard if available
        if 'standard' in section_config:
            content += f'''
            <div class="standard-box">
                <strong>STANDARD:</strong> {section_config['standard']}
            </div>
            '''
        
        # Add instructions if available
        if 'instructions' in section_config:
            content += f'''
            <div class="standard-box">
                <strong>Instructions:</strong> {section_config['instructions']}
            </div>
            '''
        
        # Add scoring legend
        if 'scoring' in section_config or section_type in ['conditional_questions', 'register_checklist']:
            content += '''
            <div class="scoring-box">
                <h3>Scoring:</h3>
                <ul>
                    <li><span class="red-dot">● Red</span> = Critical issue (Q1 or Q2 = No)</li>
                    <li><span class="yellow-dot">● Yellow</span> = Needs improvement (Q3 = No)</li>
                    <li><span class="light-green-dot">● Light Green</span> = Good (Q4 = No)</li>
                    <li><span class="dark-green-dot">● Dark Green</span> = Excellent (All Yes)</li>
                </ul>
            </div>
            '''
        
        # Generate content based on section type
        # Type 1: Register Checklist (with registers array)
        if section_type == 'register_checklist' and 'registers' in section_config:
            content += '<table class="data-table">'
            content += '<tr><th>Register Name</th>'
            for col in section_config.get('columns', ['Available', 'Standard versions', '90% complete']):
                content += f'<th>{col}</th>'
            content += '</tr>'
            
            for register in section_config['registers']:
                reg_id = register['id']
                reg_name = register['name']
                content += f'<tr><td>{reg_name}</td>'
                for column in section_config.get('columns', ['Available', 'Standard versions', '90% complete']):
                    indicator_id = f"{reg_id}_{column.lower().replace('%', '_pct').replace(' ', '_')}"
                    score = scores.get(indicator_id, '')
                    score_display = str(score) if score else 'N/A'
                    
                    # Color code the cell
                    cell_class = ''
                    if str(score).lower() in ['yes', 'y', '4']:
                        cell_class = ' style="background: #006400; color: white; font-weight: bold;"'
                    elif str(score).lower() in ['no', 'n', '1']:
                        cell_class = ' style="background: #DC3545; color: white; font-weight: bold;"'
                    elif str(score) == '2':
                        cell_class = ' style="background: #FFC107; font-weight: bold;"'
                    elif str(score) == '3':
                        cell_class = ' style="background: #90EE90; font-weight: bold;"'
                    
                    content += f'<td{cell_class}>{score_display}</td>'
                content += '</tr>'
            content += '</table>'
        
        # Type 2: Quality Matrix
        elif section_type == 'quality_matrix' and 'services' in section_config:
            content += '<table class="data-table">'
            content += '<tr><th>Service/Indicator</th><th>Response</th></tr>'
            for service in section_config['services']:
                score = scores.get(service['id'], '')
                score_display = str(score) if score else 'N/A'
                
                cell_class = ''
                if str(score).lower() in ['yes', 'y']:
                    cell_class = ' style="background: #006400; color: white; font-weight: bold;"'
                elif str(score).lower() in ['no', 'n']:
                    cell_class = ' style="background: #DC3545; color: white; font-weight: bold;"'
                
                content += f'<tr><td>{service.get("name", service["id"])}</td><td{cell_class}>{score_display}</td></tr>'
            content += '</table>'
        
        # Type 3: Questions array (check before indicators, as some sections have both)
        elif 'questions' in section_config:
            for q in section_config['questions']:
                q_id = q['id']
                q_text = q.get('text', q_id)
                
                content += f'''
                <div class="question">
                    <div class="question-label">{q_text}</div>
                '''
                
                # Handle multi_yes_no with items
                if 'items' in q and q.get('type') == 'multi_yes_no':
                    content += '<div style="margin-left: 20px;">'
                    for item in q['items']:
                        item_id = f"{q_id}_{item.lower().replace(' ', '_').replace('/', '_').replace('.', '')}"
                        score = scores.get(item_id, '')
                        
                        yes_class = 'selected' if str(score).lower() in ['yes', 'y'] else ''
                        no_class = 'selected' if str(score).lower() in ['no', 'n'] else ''
                        
                        content += f'''
                        <div style="margin: 8px 0;">
                            <strong>{item}:</strong>
                            <div class="radio-options">
                                <div class="radio-option {yes_class}">YES</div>
                                <div class="radio-option {no_class}">NO</div>
                            </div>
                        </div>
                        '''
                    content += '</div>'
                
                # Handle data_entry_table with fields
                elif 'fields' in q:
                    content += '<table class="data-table" style="margin-top: 10px;">'
                    content += '<tr><th>Field</th><th>Value</th></tr>'
                    for field in q['fields']:
                        field_id = field['id']
                        field_name = field.get('name', field.get('label', field_id))
                        score = scores.get(field_id, '')
                        content += f'<tr><td>{field_name}</td><td>{score if score else "N/A"}</td></tr>'
                    content += '</table>'
                
                # Regular question with radio buttons
                else:
                    score = scores.get(q_id, '')
                    yes_class = 'selected' if str(score).lower() in ['yes', 'y'] else ''
                    no_class = 'selected' if str(score).lower() in ['no', 'n'] else ''
                    
                    content += f'''
                    <div class="radio-options">
                        <div class="radio-option {yes_class}">YES</div>
                        <div class="radio-option {no_class}">NO</div>
                    </div>
                    '''
                
                # Add comment if exists
                if comments.get(q_id):
                    content += f'<div class="comments-box">Comment: {comments.get(q_id)}</div>'
                
                content += '</div>'
        
        # Type 4: Indicators array (standard assessment sections)
        elif 'indicators' in section_config:
            content += '<table class="data-table">'
            content += '<tr><th>Indicator</th><th>Score</th><th>Comments</th></tr>'
            for indicator in section_config['indicators']:
                ind_id = indicator['id']
                ind_name = indicator.get('name', ind_id)
                score = scores.get(ind_id, '')
                comment = comments.get(ind_id, '')
                
                # Color code based on score (1-5 scale)
                cell_class = ''
                score_display = str(score) if score else 'N/A'
                
                if str(score) in ['5', '4']:
                    cell_class = ' style="background-color: #006400; color: white; font-weight: bold;"'
                elif str(score) == '1':
                    cell_class = ' style="background-color: #DC3545; color: white; font-weight: bold;"'
                elif str(score) == '2':
                    cell_class = ' style="background-color: #FFC107; font-weight: bold;"'
                elif str(score) == '3':
                    cell_class = ' style="background-color: #90EE90; font-weight: bold;"'
                
                content += f'<tr><td>{ind_name}</td><td{cell_class}>{score_display}</td><td>{comment if comment else ""}</td></tr>'
            content += '</table>'
        
        # Fallback: If no content was generated, show all scores as a table
        if not content or content.strip() == '':
            # Display all scores from this section
            section_scores = {k: v for k, v in scores.items() if k.startswith(section_key + '_') or k == section_key}
            if section_scores:
                content += '<table class="data-table">'
                content += '<tr><th>Indicator/Question</th><th>Score</th><th>Comments</th></tr>'
                for key, score in section_scores.items():
                    if key == section_key:
                        continue  # Skip section aggregate score
                    comment = comments.get(key, '')
                    score_display = str(score) if score else 'N/A'
                    content += f'<tr><td>{key}</td><td>{score_display}</td><td>{comment if comment else ""}</td></tr>'
                content += '</table>'
            else:
                content += '<p>No assessment data available for this section.</p>'
        
        # Add section score
        section_score = scores.get(section_key, '')
        if section_score and section_score not in ['', 'N/A', None]:
            score_class = 'score-dark-green'
            score_label = 'Excellent'
            try:
                score_val = int(float(section_score))
                if score_val == 1:
                    score_class = 'score-red'
                    score_label = 'Critical'
                elif score_val == 2:
                    score_class = 'score-yellow'
                    score_label = 'Needs Improvement'
                elif score_val == 3:
                    score_class = 'score-light-green'
                    score_label = 'Good'
                else:
                    score_class = 'score-dark-green'
                    score_label = 'Excellent'
            except:
                pass
            
            content += f'''
            <div class="section-score">
                <h3>Section Score</h3>
                <span class="score-badge {score_class}">{score_label}: {section_score} / 4</span>
            </div>
            '''
        
        return content
    
    @staticmethod
    @log_function_call
    def create_registration_pdf(participants: List[Dict]) -> Tuple[str, str]:
        """Generate PDF of participant registration (screenshot-like replica)
        
        Args:
            participants: List of participant dictionaries
            
        Returns:
            Tuple of (filepath, filename)
        """
        try:
            from xhtml2pdf import pisa
            import io
            
            # Create PDF file
            temp_dir = tempfile.gettempdir()
            filename = f'participant_registration_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
            filepath = os.path.join(temp_dir, filename)
            
            # Build HTML content that matches registration screen
            html_content = PDFGenerator._build_registration_html(participants)
            
            # Generate PDF using xhtml2pdf (reliable and compatible)
            with open(filepath, 'wb') as pdf_file:
                pisa_status = pisa.CreatePDF(
                    html_content.encode('utf-8'),
                    dest=pdf_file,
                    encoding='utf-8'
                )
                
                if pisa_status.err:
                    raise Exception(f"PDF generation had {pisa_status.err} errors")
            
            # Verify file was created
            if not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
                raise Exception("PDF file was not created or is empty")
            
            return filepath, filename
            
        except Exception as e:
            raise Exception(f"Failed to create registration PDF: {str(e)}")
    
    @staticmethod
    def _build_registration_html(participants: List[Dict]) -> str:
        """Build HTML for registration PDF that matches the frontend appearance"""
        
        # Build participants table rows with alternating colors
        rows_html = ''
        for idx, p in enumerate(participants, 1):
            campaign_day = f"Day {p.get('campaignDay')}" if p.get('campaignDay') else "N/A"
            row_class = 'class="alt-row"' if idx % 2 == 0 else ''
            rows_html += f'''
            <tr {row_class}>
                <td style="text-align: center;">{idx}</td>
                <td>{p.get('fullName', 'N/A')}</td>
                <td>{p.get('cadre', 'N/A')}</td>
                <td>{p.get('facilityName', 'N/A')}</td>
                <td>{p.get('district', 'N/A')}</td>
                <td>{p.get('mobileNumber', 'N/A')}</td>
                <td>{p.get('registrationDate', 'N/A')}</td>
                <td>{campaign_day}</td>
                <td>{p.get('mobileMoneyProvider', 'N/A')}</td>
                <td>{p.get('mobileMoneyNumber', 'N/A')}</td>
                <td>{p.get('mobileMoneyName', 'N/A')}</td>
            </tr>
            '''
        
        html = f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Participant Registration</title>
    <style>
        @page {{
            size: A4 landscape;
            margin: 0.75cm;
        }}
        body {{
            font-family: Arial, Helvetica, sans-serif;
            font-size: 9pt;
            line-height: 1.4;
            color: #333;
            margin: 0;
            padding: 0;
        }}
        .header {{
            background-color: #667eea;
            color: white;
            padding: 15px 20px;
            margin-bottom: 15px;
            text-align: center;
        }}
        .header h1 {{
            margin: 0 0 5px 0;
            font-size: 20pt;
            font-weight: bold;
        }}
        .header p {{
            margin: 0;
            font-size: 11pt;
        }}
        .info-box {{
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
            padding: 10px;
            margin-bottom: 15px;
        }}
        .info-box table {{
            width: 100%;
            border: none;
        }}
        .info-box td {{
            border: none;
            padding: 5px;
        }}
        .count {{
            font-size: 14pt;
            font-weight: bold;
            color: #667eea;
        }}
        .date {{
            font-size: 9pt;
            color: #666;
            text-align: right;
        }}
        table.data-table {{
            width: 100%;
            border-collapse: collapse;
            background: white;
        }}
        table.data-table th {{
            background-color: #667eea;
            color: white;
            padding: 8px 6px;
            text-align: left;
            font-weight: bold;
            font-size: 8.5pt;
            border: 1px solid #5a67d8;
        }}
        table.data-table td {{
            padding: 6px;
            border: 1px solid #dee2e6;
            font-size: 8.5pt;
        }}
        .alt-row {{
            background-color: #f8f9fa;
        }}
        .footer {{
            margin-top: 15px;
            padding-top: 10px;
            border-top: 2px solid #667eea;
            text-align: center;
            font-size: 8pt;
            color: #666;
        }}
        .footer strong {{
            color: #667eea;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Participant Registration Report</h1>
        <p>CHAI - Clinton Health Access Initiative</p>
    </div>
    
    <div class="info-box">
        <table>
            <tr>
                <td><span class="count">Total Participants: {len(participants)}</span></td>
                <td class="date">Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}</td>
            </tr>
        </table>
    </div>
    
    <table class="data-table">
        <thead>
            <tr>
                <th style="width: 3%;">#</th>
                <th style="width: 12%;">Full Name</th>
                <th style="width: 10%;">Cadre</th>
                <th style="width: 12%;">Facility</th>
                <th style="width: 10%;">District</th>
                <th style="width: 9%;">Mobile</th>
                <th style="width: 8%;">Reg. Date</th>
                <th style="width: 6%;">Campaign</th>
                <th style="width: 8%;">MM Provider</th>
                <th style="width: 9%;">MM Number</th>
                <th style="width: 13%;">MM Name</th>
            </tr>
        </thead>
        <tbody>
            {rows_html}
        </tbody>
    </table>
    
    <div class="footer">
        <p><strong>CHAI Uganda</strong> | Health Worker Registration and Tracking System</p>
        <p>This is an automatically generated report. For inquiries, contact the CHAI team.</p>
    </div>
</body>
</html>
'''
        return html